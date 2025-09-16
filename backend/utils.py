import os
import subprocess
from docx import Document
from openpyxl import Workbook, load_workbook


def compress_pdf(input_path, output_path, target_kb):
    """
    Compress PDF using Ghostscript. Chooses compression level
    based on the target size (approximate).
    """
    if not os.path.exists(input_path):
        raise FileNotFoundError(f"PDF file not found: {input_path}")

    # Decide quality from target size
    if target_kb <= 500:
        quality = "screen"     # smallest
    elif target_kb <= 1500:
        quality = "ebook"      # medium
    else:
        quality = "printer"    # higher quality

    gs_cmd = [
        "gs",
        "-sDEVICE=pdfwrite",
        "-dCompatibilityLevel=1.4",
        f"-dPDFSETTINGS=/{quality}",
        "-dNOPAUSE",
        "-dQUIET",
        "-dBATCH",
        f"-sOutputFile={output_path}",
        input_path,
    ]

    try:
        subprocess.run(gs_cmd, check=True)
        print(f"Compressed PDF saved to {output_path}")
    except FileNotFoundError:
        raise EnvironmentError(
            "Ghostscript (gs) not found. Install from https://www.ghostscript.com/"
        )


def word_to_excel(word_path, excel_path):
    if not os.path.exists(word_path):
        raise FileNotFoundError(f"Word file not found: {word_path}")

    doc = Document(word_path)
    wb = Workbook()
    ws = wb.active
    if ws is None:
        ws = wb.create_sheet("Sheet1")

    for i, para in enumerate(doc.paragraphs, 1):
        text = para.text.strip()
        if text:
            ws[f"A{i}"] = text

    wb.save(excel_path)
    print(f"Excel file saved to {excel_path}")


def excel_to_word(excel_path, word_path):
    if not os.path.exists(excel_path):
        raise FileNotFoundError(f"Excel file not found: {excel_path}")

    wb = load_workbook(excel_path)
    ws = wb.active
    if ws is None:
        raise ValueError("No active worksheet found in the Excel file.")

    doc = Document()
    for row in ws.iter_rows(values_only=True):
        if row:
            line = " | ".join(str(c) for c in row if c is not None)
            doc.add_paragraph(line)

    doc.save(word_path)
    print(f"Word file saved to {word_path}")
