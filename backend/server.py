# app.py
from flask import Flask, request, jsonify, send_file, Response, make_response
from flask_cors import CORS
from werkzeug.utils import secure_filename
import os
import io
import fitz  # PyMuPDF
from PIL import Image
from docx import Document
from openpyxl import Workbook, load_workbook
import threading
import time

app = Flask(__name__)
CORS(app)
UPLOAD_FOLDER = "uploads"
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

ALLOWED_PDF = {".pdf"}
ALLOWED_WORD = {".doc", ".docx"}
ALLOWED_IMAGE = {".jpg", ".jpeg", ".png", ".webp", ".tiff", ".bmp"}
ALLOWED_EXCEL = {".xls", ".xlsx"}

MAX_FILE_AGE_SECONDS = 3600  # 1 hour for cleanup

def cleanup_files():
    """Deletes files in the upload folder older than MAX_FILE_AGE_SECONDS."""
    while True:
        now = time.time()
        for filename in os.listdir(app.config['UPLOAD_FOLDER']):
            filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            if os.path.isfile(filepath):
                file_age = now - os.path.getmtime(filepath)
                if file_age > MAX_FILE_AGE_SECONDS:
                    print(f"Deleting old file: {filename}")
                    try:
                        os.remove(filepath)
                    except OSError as e:
                        print(f"Error deleting file {filepath}: {e}")
        time.sleep(600)  # Check every 10 minutes

# Start the cleanup thread
cleanup_thread = threading.Thread(target=cleanup_files)
cleanup_thread.daemon = True
cleanup_thread.start()

def save_uploaded_file(uploaded):
    filename = secure_filename(uploaded.filename or "uploaded.bin")
    path = os.path.join(UPLOAD_FOLDER, filename)
    uploaded.save(path)
    return filename, path

def error_response(msg: str, status: int = 500) -> Response:
    return make_response(jsonify({"error": msg}), status)

# -------------------- Compress PDF -------------------- #
@app.route("/compress-pdf", methods=["POST"])
def compress_pdf():
    try:
        uploaded = request.files.get("file")
        if not uploaded or not uploaded.filename:
            return error_response("No file uploaded")

        filename = secure_filename(uploaded.filename)
        ext = os.path.splitext(filename)[1].lower()
        if ext not in ALLOWED_PDF:
            return error_response("Only PDF files are allowed")

        input_path = os.path.join(UPLOAD_FOLDER, filename)
        uploaded.save(input_path)

        out_filename = f"compressed_{int(time.time())}_{filename}"
        output_path = os.path.join(UPLOAD_FOLDER, out_filename)

        quality = 50  # fixed compression quality

        with fitz.open(input_path) as doc:
            for page in doc:
                for img_info in page.get_images(full=True):
                    xref = img_info[0]
                    try:
                        img_dict = doc.extract_image(xref)
                        img_bytes = img_dict.get("image")
                        if not img_bytes:
                            continue
                        image = Image.open(io.BytesIO(img_bytes)).convert("RGB")
                        buf = io.BytesIO()
                        image.save(buf, format="JPEG", quality=quality, optimize=True)
                        doc.update_stream(xref, buf.getvalue())
                    except Exception:
                        continue
            doc.save(output_path, deflate=True, garbage=4, clean=True)

        size_kb = round(os.path.getsize(output_path) / 1024, 2)
        return jsonify({
            "status": "success",
            "download_url": f"/download/{out_filename}",
            "compressed_size_kb": size_kb
        })
    except Exception as e:
        return error_response(str(e))

# -------------------- PDF Editor -------------------- #
@app.route("/pdf-editor", methods=["POST"])
def pdf_editor() -> Response:
    """
    PDF editor actions:
      - action=rotate & angle=<int> : rotate all pages by angle degrees
      - action=split & pages=1,3-5 : keep only specified pages (1-based)
      - action=merge with multiple 'files' form fields: merges uploaded PDFs in order
    """
    try:
        action = request.form.get("action", "").lower()
        if action not in {"rotate", "split", "merge"}:
            return error_response("action must be one of: rotate, split, merge", 400)

        if action == "merge":
            files = request.files.getlist("files")
            if not files:
                return error_response("No files provided for merge", 400)

            merged = fitz.open()
            for f in files:
                doc = fitz.open(stream=f.read(), filetype="pdf")
                merged.insert_pdf(doc)  # type: ignore
                doc.close()

            out_name = f"merged_{len(files)}_{int(time.time())}.pdf"
            out_path = os.path.join(UPLOAD_FOLDER, out_name)
            merged.save(out_path)
            merged.close()
            return jsonify({"status": "success", "download_url": f"/download/{out_name}"})

        # for rotate and split: expect single 'file'
        uploaded = request.files.get("file")
        if not uploaded or not uploaded.filename:
            return error_response("No file uploaded", 400)

        filename, input_path = save_uploaded_file(uploaded)
        out_basename = f"edited_{int(time.time())}_{filename}"
        out_path = os.path.join(UPLOAD_FOLDER, out_basename)

        if action == "rotate":
            try:
                angle = int(request.form.get("angle", 90))
            except ValueError:
                return error_response("angle must be an integer", 400)

            with fitz.open(input_path) as doc:
                for p in range(len(doc)):
                    page = doc[p]
                    page.set_rotation((page.rotation + angle) % 360)  # type: ignore
                doc.save(out_path)
            return jsonify({"status": "success", "download_url": f"/download/{out_basename}"})

        if action == "split":
            pages_spec = request.form.get("pages", "")
            if not pages_spec:
                return error_response("Provide pages parameter, e.g. pages=1,3-5", 400)

            # parse pages_spec like "1,3-5" -> zero-based indices [0,2,3,4]
            pages_to_keep = []
            for part in pages_spec.split(","):
                part = part.strip()
                if "-" in part:
                    a, b = part.split("-", 1)
                    a_i = int(a.strip()) - 1
                    b_i = int(b.strip()) - 1
                    pages_to_keep.extend(list(range(a_i, b_i + 1)))
                else:
                    pages_to_keep.append(int(part) - 1)

            with fitz.open(input_path) as doc:
                newdoc = fitz.open()
                for idx in pages_to_keep:
                    if 0 <= idx < len(doc):
                        newdoc.insert_pdf(doc, from_page=idx, to_page=idx)  # type: ignore
                newdoc.save(out_path)
                newdoc.close()
            return jsonify({"status": "success", "download_url": f"/download/{out_basename}"})

        # If somehow no action matched
        return error_response("Unknown error occurred", 500)

    except Exception as e:
        return error_response(str(e), 500)

# -------------------- Word -> Excel -------------------- #
@app.route("/word-to-excel", methods=["POST"])
def word_to_excel():
    """
    Extracts tables in a Word (.docx) into a single Excel workbook.
    """
    try:
        uploaded = request.files.get("file")
        if not uploaded or not uploaded.filename:
            return error_response("No file uploaded", 400)
        filename, input_path = save_uploaded_file(uploaded)
        out_name = f"{os.path.splitext(filename)[0]}_{int(time.time())}.xlsx"
        out_path = os.path.join(UPLOAD_FOLDER, out_name)

        try:
            doc = Document(input_path)
            wb = Workbook()
            ws = wb.active
            row_index = 1
            for table in doc.tables:
                for row in table.rows:
                    col_index = 1
                    for cell in row.cells:
                        ws.cell(row=row_index, column=col_index, value=cell.text.strip())  # type: ignore
                        col_index += 1
                    row_index += 1
                row_index += 1  # blank line between tables
            wb.save(out_path)
        except Exception as e:
            return error_response(f"Conversion failed: {e}")

        return jsonify({"status": "success", "download_url": f"/download/{out_name}"})
    except Exception as e:
        return error_response(str(e))

# -------------------- Excel -> Word -------------------- #
@app.route("/excel-to-word", methods=["POST"])
def excel_to_word():
    """
    Reads the first sheet of an Excel file and writes its contents into a Word document as paragraphs.
    """
    try:
        uploaded = request.files.get("file")
        if not uploaded or not uploaded.filename:
            return error_response("No file uploaded", 400)

        filename, input_path = save_uploaded_file(uploaded)
        out_name = f"{os.path.splitext(filename)[0]}_{int(time.time())}.docx"
        out_path = os.path.join(UPLOAD_FOLDER, out_name)

        try:
            # Open workbook safely in read-only mode
            wb = load_workbook(input_path, read_only=True, data_only=True)
            if not wb.sheetnames:
                return error_response("Excel file has no sheets", 400)

            ws = wb.active

            # ADD THIS CHECK: Ensure the active sheet exists
            if ws is None:
                return error_response("The active sheet is empty or does not exist", 400)

            rows = list(ws.iter_rows(values_only=True))

            if not rows:
                return error_response("Excel sheet is empty", 400)

            # Create a Word document
            doc = Document()

            # Add rows as tab-separated lines
            for row in rows:
                line = "\t".join("" if v is None else str(v) for v in row)
                doc.add_paragraph(line)

            doc.save(out_path)
        except Exception as e:
            return error_response(f"Conversion failed: {e}")

        return jsonify({"status": "success", "download_url": f"/download/{out_name}"})
    except Exception as e:
        return error_response(str(e))

# -------------------- Image Compress -------------------- #
@app.route("/image-compress", methods=["POST"])
def image_compress():
    """
    Compress an uploaded image (file) and returns compressed image.
    Optional form params:
      - quality (int 1-95), default 50
      - max_width, max_height to resize keeping aspect ratio
    """
    try:
        uploaded = request.files.get("file")
        if not uploaded or not uploaded.filename:
            return error_response("No image uploaded", 400)

        filename = secure_filename(uploaded.filename)
        ext = os.path.splitext(filename)[1].lower()
        if ext not in ALLOWED_IMAGE:
            return error_response("Unsupported image type", 400)

        quality = int(request.form.get("quality", 50))
        max_w = request.form.get("max_width")
        max_h = request.form.get("max_height")
        max_w = int(max_w) if max_w else None
        max_h = int(max_h) if max_h else None

        image = Image.open(io.BytesIO(uploaded.read())).convert("RGB")
        # optional resize while keeping ratio
        if max_w or max_h:
            w, h = image.size
            ratio = min((max_w or w) / w, (max_h or h) / h)
            if ratio < 1:
                new_size = (int(w * ratio), int(h * ratio))
                image = image.resize(new_size, Image.Resampling.LANCZOS)

        out_buf = io.BytesIO()
        image.save(out_buf, format="JPEG", quality=quality, optimize=True)
        out_buf.seek(0)

        out_name = f"compressed_{int(time.time())}_{os.path.splitext(filename)[0]}.jpg"
        save_path = os.path.join(UPLOAD_FOLDER, out_name)
        with open(save_path, "wb") as f:
            f.write(out_buf.getvalue())

        return jsonify({"status": "success", "download_url": f"/download/{out_name}"})
    except Exception as e:
        return error_response(str(e))

# -------------------- PDF -> Word -------------------- #
@app.route("/pdf-to-word", methods=["POST"])
def pdf_to_word():
    """
    Extracts text from PDF pages and writes into a .docx. (Does not yet extract images)
    """
    try:
        uploaded = request.files.get("file")
        if not uploaded or not uploaded.filename:
            return error_response("No PDF uploaded", 400)

        filename, input_path = save_uploaded_file(uploaded)
        out_name = f"{os.path.splitext(filename)[0]}_{int(time.time())}.docx"
        out_path = os.path.join(UPLOAD_FOLDER, out_name)

        try:
            docx_doc = Document()
            with fitz.open(input_path) as pdf:  # type: ignore
                for page in pdf:
                    text = page.get_text()  # type: ignore
                    if text and text.strip():
                        # split to paragraphs by double newline for nicer docx
                        for para in text.split("\n\n"):
                            docx_doc.add_paragraph(para.strip())
                    docx_doc.add_page_break()
            docx_doc.save(out_path)
        except Exception as e:
            return error_response(f"Conversion failed: {e}")

        return jsonify({"status": "success", "download_url": f"/download/{out_name}"})
    except Exception as e:
        return error_response(str(e))

# -------------------- Download -------------------- #
@app.route("/download/<path:filename>", methods=["GET"])
def download_file(filename):
    try:
        file_path = os.path.join(UPLOAD_FOLDER, filename)
        if not os.path.exists(file_path):
            return error_response("File not found", 404)
        # choose mimetype for common types
        ext = os.path.splitext(filename)[1].lower()
        if ext == ".pdf":
            mimetype = "application/pdf"
        elif ext in {".docx", ".doc"}:
            mimetype = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        elif ext in {".xlsx", ".xls"}:
            mimetype = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        elif ext in {".jpg", ".jpeg"}:
            mimetype = "image/jpeg"
        else:
            mimetype = None
        return send_file(file_path, mimetype=mimetype, as_attachment=True)
    except Exception as e:
        return error_response(str(e))


if __name__ == "__main__":
    app.run(host="0.0.0.0", port=5000)